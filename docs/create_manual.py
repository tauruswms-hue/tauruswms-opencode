from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# ── Title ──
title = doc.add_heading('Taurus WMS — Configuración de Engine de Base de Datos', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph(
    'Guía paso a paso para configurar el motor de base de datos '
    'en Taurus WMS. Soporta MySQL, PostgreSQL, SQL Server y SQLite.'
)

# ── Engines soportados ──
doc.add_heading('Engines soportados', level=1)

table = doc.add_table(rows=5, cols=4)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ['Engine', 'DB_ENGINE', 'Driver', 'Puerto default']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
    for p in table.rows[0].cells[i].paragraphs:
        for run in p.runs:
            run.bold = True

data = [
    ['MySQL',       'mysql',       'pymysql',    '3306'],
    ['PostgreSQL',  'postgresql',  'psycopg2',   '5432'],
    ['SQL Server',  'sqlserver',   'pymssql',    '1433'],
    ['SQLite',      'sqlite',      'sqlite3',    'N/A'],
]
for r, row_data in enumerate(data, 1):
    for c, val in enumerate(row_data):
        table.rows[r].cells[c].text = val

# ── Paso 1 ──
doc.add_heading('Paso 1: Instalar el driver', level=1)

doc.add_paragraph('MySQL (ya incluido en requirements.txt):')
doc.add_paragraph('pip install pymysql', style='List Bullet')

doc.add_paragraph('PostgreSQL:')
doc.add_paragraph('pip install psycopg2-binary', style='List Bullet')

doc.add_paragraph('SQL Server:')
doc.add_paragraph('pip install pymssql', style='List Bullet')

doc.add_paragraph('SQLite:')
doc.add_paragraph('No necesita instalación — viene con Python.', style='List Bullet')

# ── Paso 2 ──
doc.add_heading('Paso 2: Crear la base de datos', level=1)

doc.add_paragraph(
    'En el motor correspondiente, crear la base de datos taurus_wms '
    'y ejecutar las migraciones que se encuentran en migrations/. '
    'Si partís de cero, primero creá las tablas con crear_tablas.py '
    '(solo funciona con MySQL).'
)

# ── Paso 3 ──
doc.add_heading('Paso 3: Configurar la conexión', level=1)

doc.add_heading('Opción A — Tabla configuracion (recomendado)', level=2)

doc.add_paragraph(
    'Se configura desde la admin UI o ejecutando el siguiente SQL '
    'contra la base taurus_admin:'
)

sql_block = doc.add_paragraph()
sql_block.style = 'No Spacing'
run = sql_block.add_run(
    "INSERT INTO configuracion (clave, valor, descripcion) VALUES\n"
    "('DB_ENGINE',   'mysql',         'Motor: mysql, postgresql, sqlite, sqlserver'),\n"
    "('DB_HOST',     'localhost',     'Host de la base de datos'),\n"
    "('DB_PORT',     '3306',          'Puerto'),\n"
    "('DB_NAME',     'taurus_wms',    'Nombre de la base de datos'),\n"
    "('DB_USER',     'taurus',        'Usuario'),\n"
    "('DB_PASSWORD', 'tu_password',   'Contraseña'),\n"
    "('DB_CHAR_SET', 'utf8mb4',       'Charset')\n"
    "ON DUPLICATE KEY UPDATE valor = VALUES(valor);"
)
run.font.name = 'Consolas'
run.font.size = Pt(9)

doc.add_heading('Opción B — Archivo .env (override para desarrollo)', level=2)

env_block = doc.add_paragraph()
env_block.style = 'No Spacing'
run = env_block.add_run(
    "DB_ENGINE=postgresql\n"
    "DB_HOST=localhost\n"
    "DB_PORT=5432\n"
    "DB_NAME=taurus_wms\n"
    "DB_USER=taurus\n"
    "DB_PASSWORD=secret\n"
    "DB_CHAR_SET=UTF8"
)
run.font.name = 'Consolas'
run.font.size = Pt(9)

# ── Paso 4 ──
doc.add_heading('Paso 4: Admin DB (siempre MySQL)', level=1)

doc.add_paragraph(
    'La base de datos taurus_admin (usuarios, tenants, configuración) '
    'siempre utiliza MySQL. No hay forma de cambiar esto. '
    'Se configura en el archivo .env:'
)

admin_block = doc.add_paragraph()
admin_block.style = 'No Spacing'
run = admin_block.add_run(
    "DB_ADMIN_HOST=localhost\n"
    "DB_ADMIN_PORT=3306\n"
    "DB_ADMIN_NAME=taurus_admin\n"
    "DB_ADMIN_USER=taurus_admin\n"
    "DB_ADMIN_PASSWORD=Taurus_2001"
)
run.font.name = 'Consolas'
run.font.size = Pt(9)

# ── Paso 5 ──
doc.add_heading('Paso 5: Verificar', level=1)

doc.add_paragraph('Ejecutar la aplicación:')
doc.add_paragraph('python app.py', style='List Bullet')

doc.add_paragraph(
    'Si el driver no está instalado, se mostrará un ImportError '
    'con el comando pip install correcto.'
)

# ── Flujo de resolución ──
doc.add_heading('Flujo de resolución de configuración', level=1)

doc.add_paragraph(
    'El motor se resuelve en el siguiente orden de prioridad:'
)

doc.add_paragraph('1. Archivo .env → variable DB_ENGINE', style='List Number')
doc.add_paragraph('   Si existe → se usa.', style='List Bullet')
doc.add_paragraph('   Si no existe → paso 2.', style='List Bullet')

doc.add_paragraph('2. Tabla configuracion → registro DB_ENGINE', style='List Number')
doc.add_paragraph('   Si existe → se usa.', style='List Bullet')
doc.add_paragraph('   Si no existe → default mysql.', style='List Bullet')

doc.add_paragraph(
    'Las credenciales (DB_HOST, DB_PORT, DB_NAME, DB_USER, DB_PASSWORD) '
    'se resolved同样的 prioridad: .env primero, tabla configuracion después.'
)

# ── Notas importantes ──
doc.add_heading('Notas importantes', level=1)

doc.add_paragraph(
    'PostgreSQL: Requiere psycopg2-binary. Los cursores devuelven '
    'dicts (RealDictCursor), igual que MySQL.',
    style='List Bullet'
)
doc.add_paragraph(
    'SQL Server: Requiere pymssql. Usa MERGE para upserts, '
    'SCOPE_IDENTITY() para obtener el último ID, y OFFSET/FETCH NEXT '
    'en vez de LIMIT.',
    style='List Bullet'
)
doc.add_paragraph(
    'SQLite: No necesita driver externo. No soporta concurrencia '
    'ni múltiples conexiones simultáneas. Ideal para testing/desarrollo.',
    style='List Bullet'
)
doc.add_paragraph(
    'Cache: La configuración se cachea en memoria. Para recargar, '
    'llamar clear_config_cache() desde módulos.db_config.',
    style='List Bullet'
)
doc.add_paragraph(
    'Cambiar engine: Modificar DB_ENGINE en .env o en la tabla '
    'configuracion y reiniciar la aplicación.',
    style='List Bullet'
)

# ── Guardar ──
output_path = 'docs/configuracion_engineBD.docx'
doc.save(output_path)
print(f'Documento guardado en: {output_path}')
